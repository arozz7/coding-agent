"""Document reader tool — PDF, DOCX, XLSX, CSV.

All readers return a consistent dict:
  {"success": bool, "type": str, "text"|"data": ..., "truncated": bool, ...}

Text is capped at _MAX_CHARS; tabular data at _MAX_ROWS rows.
Each reader degrades gracefully if its optional dependency is absent.
"""

import csv
from pathlib import Path
from typing import Any, Dict, List, Optional

import structlog

logger = structlog.get_logger()

_MAX_CHARS = 10_000   # max characters of extracted text per document
_MAX_ROWS = 500       # max rows returned from spreadsheets / CSVs

_SUPPORTED = {".pdf", ".docx", ".doc", ".xlsx", ".xls", ".csv", ".tsv"}


class DocumentTool:
    """Read PDF, Word, Excel, and CSV/TSV files into plain text or structured data."""

    def __init__(self):
        self.logger = logger.bind(component="document_tool")

    # ------------------------------------------------------------------
    # Public interface
    # ------------------------------------------------------------------

    def read(self, path: str) -> Dict[str, Any]:
        """Auto-detect document type by extension and read it."""
        p = Path(path)
        if not p.exists():
            return {"success": False, "error": f"File not found: {path}"}

        ext = p.suffix.lower()
        dispatch = {
            ".pdf":  self.read_pdf,
            ".docx": self.read_docx,
            ".doc":  self.read_docx,
            ".xlsx": self.read_xlsx,
            ".xls":  self.read_xlsx,
            ".csv":  self.read_csv,
            ".tsv":  self.read_csv,
        }
        reader = dispatch.get(ext)
        if not reader:
            return {
                "success": False,
                "error": f"Unsupported format '{ext}'. Supported: {sorted(_SUPPORTED)}",
            }
        return reader(path)

    # ------------------------------------------------------------------
    # Per-format readers
    # ------------------------------------------------------------------

    def read_pdf(self, path: str) -> Dict[str, Any]:
        """Extract text from a PDF using pypdf."""
        try:
            import pypdf  # type: ignore
        except ImportError:
            return {"success": False, "error": "pypdf not installed — run: pip install pypdf"}

        try:
            reader = pypdf.PdfReader(path)
            total_pages = len(reader.pages)
            parts: List[str] = []
            chars = 0

            for i, page in enumerate(reader.pages):
                page_text = page.extract_text() or ""
                parts.append(f"[Page {i + 1}]\n{page_text}")
                chars += len(page_text)
                if chars >= _MAX_CHARS:
                    parts.append(f"[truncated — {total_pages - i - 1} more pages]")
                    break

            text = "\n\n".join(parts)
            return {
                "success": True,
                "path": path,
                "type": "pdf",
                "total_pages": total_pages,
                "text": text[:_MAX_CHARS],
                "truncated": chars >= _MAX_CHARS,
            }
        except Exception as e:
            self.logger.error("pdf_read_failed", path=path, error=str(e))
            return {"success": False, "error": str(e)}

    def read_docx(self, path: str) -> Dict[str, Any]:
        """Extract paragraphs from a Word document using python-docx."""
        try:
            import docx  # type: ignore
        except ImportError:
            return {"success": False, "error": "python-docx not installed — run: pip install python-docx"}

        try:
            doc = docx.Document(path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

            # Also extract table cell text
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        paragraphs.append(row_text)

            text = "\n".join(paragraphs)
            return {
                "success": True,
                "path": path,
                "type": "docx",
                "paragraph_count": len(paragraphs),
                "text": text[:_MAX_CHARS],
                "truncated": len(text) > _MAX_CHARS,
            }
        except Exception as e:
            self.logger.error("docx_read_failed", path=path, error=str(e))
            return {"success": False, "error": str(e)}

    def read_xlsx(self, path: str) -> Dict[str, Any]:
        """Read an Excel workbook sheet-by-sheet using openpyxl."""
        try:
            import openpyxl  # type: ignore
        except ImportError:
            return {"success": False, "error": "openpyxl not installed — run: pip install openpyxl"}

        try:
            wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
            sheets: Dict[str, List[List[str]]] = {}
            total_rows = 0

            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                rows: List[List[str]] = []
                for row in ws.iter_rows(values_only=True):
                    rows.append([str(c) if c is not None else "" for c in row])
                    total_rows += 1
                    if total_rows >= _MAX_ROWS:
                        break
                sheets[sheet_name] = rows

            wb.close()
            return {
                "success": True,
                "path": path,
                "type": "xlsx",
                "sheets": list(sheets.keys()),
                "data": sheets,
                "truncated": total_rows >= _MAX_ROWS,
            }
        except Exception as e:
            self.logger.error("xlsx_read_failed", path=path, error=str(e))
            return {"success": False, "error": str(e)}

    def read_csv(self, path: str) -> Dict[str, Any]:
        """Read a CSV or TSV file using the stdlib csv module."""
        try:
            delimiter = "\t" if path.endswith(".tsv") else ","
            rows: List[List[str]] = []
            with open(path, newline="", encoding="utf-8", errors="replace") as f:
                for i, row in enumerate(csv.reader(f, delimiter=delimiter)):
                    rows.append(row)
                    if i >= _MAX_ROWS:
                        break

            return {
                "success": True,
                "path": path,
                "type": "csv",
                "rows": len(rows),
                "headers": rows[0] if rows else [],
                "data": rows,
                "truncated": len(rows) >= _MAX_ROWS,
            }
        except Exception as e:
            self.logger.error("csv_read_failed", path=path, error=str(e))
            return {"success": False, "error": str(e)}

    def supported_extensions(self) -> List[str]:
        return sorted(_SUPPORTED)


__all__ = ["DocumentTool"]
